# import os
# import json
# import re
# from flask import Flask, render_template, request, send_file, jsonify
# import pdfplumber
# import docx
# from werkzeug.utils import secure_filename
# from fpdf import FPDF
# from langchain_google_genai import ChatGoogleGenerativeAI
# from langchain_core.prompts import PromptTemplate
# from dotenv import load_dotenv

# load_dotenv()

# # Flask app setup
# app = Flask(__name__)
# app.config['UPLOAD_FOLDER'] = 'uploads/'
# app.config['RESULTS_FOLDER'] = 'results/'
# app.config['ALLOWED_EXTENSIONS'] = {'pdf', 'txt', 'docx'}

# # Ensure directories exist
# os.makedirs(app.config['UPLOAD_FOLDER'], exist_ok=True)
# os.makedirs(app.config['RESULTS_FOLDER'], exist_ok=True)

# # Initialize LangChain LLM
# llm = ChatGoogleGenerativeAI(
#     model="gemini-2.5-flash",
#     api_key=os.getenv("GOOGLE_API_KEY"),
#     temperature=0.0
# )

# # LangChain prompt template
# mcq_prompt = PromptTemplate(
#     input_variables=["context", "num_questions"],
#     template="""
# You are an AI assistant helping the user generate multiple-choice questions (MCQs) from the text below:

# Text:
# {context}

# Generate {num_questions} MCQs. For each question, provide exactly 4 options.

# Respond ONLY with the MCQs in this exact format (no other text):
# Question 1: [question text]
# A) [option A]
# B) [option B]
# C) [option C]
# D) [option D]
# Correct Answer: [A/B/C/D]

# Question 2: [question text]
# A) [option A]
# B) [option B]
# C) [option C]
# D) [option D]
# Correct Answer: [A/B/C/D]

# (Continue for all {num_questions} questions)
# """
# )

# mcq_chain = mcq_prompt | llm

# # File validation
# def allowed_file(filename):
#     return '.' in filename and filename.rsplit('.', 1)[1].lower() in app.config['ALLOWED_EXTENSIONS']

# # Text extraction
# def extract_text_from_file(file_path):
#     ext = file_path.rsplit('.', 1)[1].lower()
#     if ext == 'pdf':
#         with pdfplumber.open(file_path) as pdf:
#             return ''.join([page.extract_text() for page in pdf.pages if page.extract_text()])
#     elif ext == 'docx':
#         doc = docx.Document(file_path)
#         return ' '.join([para.text for para in doc.paragraphs])
#     elif ext == 'txt':
#         with open(file_path, 'r', encoding='utf-8') as file:
#             return file.read()
#     return None

# # MCQ generation
# def generate_mcqs_with_langchain(text, num_questions):
#     response = mcq_chain.invoke({"context": text, "num_questions": num_questions})
#     return response.content.strip()

# # Parse MCQs from text to structured format
# def parse_mcqs_to_json(mcq_text, quiz_title="Unknown Quiz"):
#     """
#     Parse MCQ text output into JSON format.
#     Expected format:
#     Question 1: [question]
#     A) [option A]
#     B) [option B]
#     C) [option C]
#     D) [option D]
#     Correct Answer: [A/B/C/D]
#     """
#     questions = []
    
#     # Split by "Question X:"
#     question_blocks = re.split(r'Question \d+:', mcq_text)
    
#     for block in question_blocks[1:]:  # Skip first empty split
#         lines = [line.strip() for line in block.strip().split('\n') if line.strip()]
        
#         if len(lines) < 6:  # Need at least question + 4 options + correct answer
#             continue
        
#         question_text = lines[0]
#         options = []
#         correct_option = -1
        
#         for i, line in enumerate(lines[1:]):
#             if line.startswith(('A)', 'a)')):
#                 options.append(line[2:].strip())
#             elif line.startswith(('B)', 'b)')):
#                 options.append(line[2:].strip())
#             elif line.startswith(('C)', 'c)')):
#                 options.append(line[2:].strip())
#             elif line.startswith(('D)', 'd)')):
#                 options.append(line[2:].strip())
#             elif 'Correct Answer:' in line:
#                 # Extract the correct answer letter
#                 answer_match = re.search(r'[A-D]', line.upper())
#                 if answer_match:
#                     answer_letter = answer_match.group()
#                     correct_option = ord(answer_letter) - ord('A')  # Convert A=0, B=1, C=2, D=3
        
#         if len(options) == 4 and correct_option >= 0:
#             questions.append({
#                 "question": question_text,
#                 "options": options,
#                 "correct_option": correct_option
#             })
    
#     return {
#         "quiz_title": quiz_title,
#         "questions": questions
#     }

# # Save MCQs to text file
# def save_mcqs_to_file(mcqs, filename):
#     path = os.path.join(app.config['RESULTS_FOLDER'], filename)
#     with open(path, 'w', encoding='utf-8') as f:
#         f.write(mcqs)
#     return path

# # Save MCQs to JSON file
# def save_mcqs_to_json(quiz_data, filename):
#     path = os.path.join(app.config['RESULTS_FOLDER'], filename)
#     with open(path, 'w', encoding='utf-8') as f:
#         json.dump(quiz_data, f, indent=2, ensure_ascii=False)
#     return path

# # Save MCQs to PDF
# def create_pdf(mcqs, filename):
#     pdf = FPDF()
#     pdf.add_page()
#     pdf.set_font("Arial", size=12)

#     for mcq in mcqs.split("## MCQ"):
#         if mcq.strip():
#             pdf.multi_cell(0, 10, mcq.strip())
#             pdf.ln(5)

#     path = os.path.join(app.config['RESULTS_FOLDER'], filename)
#     pdf.output(path)
#     return path

# # Routes
# @app.route('/')
# def index():
#     return render_template('index.html')

# @app.route('/generate', methods=['POST'])
# def generate_mcqs():
#     if 'file' not in request.files:
#         return "No file uploaded."

#     file = request.files['file']
#     if file and allowed_file(file.filename):
#         filename = secure_filename(file.filename)
#         file_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
#         file.save(file_path)

#         text = extract_text_from_file(file_path)
#         if text:
#             num_questions = int(request.form['num_questions'])
#             mcqs = generate_mcqs_with_langchain(text, num_questions)

#             # Save output
#             base_name = filename.rsplit('.', 1)[0]
#             txt_file = f"generated_mcqs_{base_name}.txt"
#             pdf_file = f"generated_mcqs_{base_name}.pdf"
#             json_file = f"generated_mcqs_{base_name}.json"
            
#             save_mcqs_to_file(mcqs, txt_file)
#             create_pdf(mcqs, pdf_file)
            
#             # Parse and save as JSON
#             quiz_data = parse_mcqs_to_json(mcqs, quiz_title=base_name)
#             save_mcqs_to_json(quiz_data, json_file)

#             return render_template('results.html', mcqs=mcqs, txt_filename=txt_file, pdf_filename=pdf_file, json_filename=json_file)

#     return "Invalid file format or upload error."

# @app.route('/generate_json', methods=['POST'])
# def generate_mcqs_json():
#     """Generate MCQs and return as JSON"""
#     if 'file' not in request.files:
#         return jsonify({"error": "No file uploaded."}), 400

#     file = request.files['file']
#     if not (file and allowed_file(file.filename)):
#         return jsonify({"error": "Invalid file format or upload error."}), 400

#     filename = secure_filename(file.filename)
#     file_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
#     file.save(file_path)

#     text = extract_text_from_file(file_path)
#     if not text:
#         return jsonify({"error": "Could not extract text from file."}), 400

#     num_questions = int(request.form.get('num_questions', 5))
#     mcqs = generate_mcqs_with_langchain(text, num_questions)
    
#     # Parse and return as JSON
#     base_name = filename.rsplit('.', 1)[0]
#     quiz_data = parse_mcqs_to_json(mcqs, quiz_title=base_name)
    
#     return jsonify(quiz_data)

# @app.route('/download/<filename>')
# def download_file(filename):
#     path = os.path.join(app.config['RESULTS_FOLDER'], filename)
#     return send_file(path, as_attachment=True)

# if __name__ == "__main__":
#     app.run(debug=False)






import os
import json
import re
import io
from flask import Flask, render_template, request, send_file, jsonify
import pdfplumber
import docx
from werkzeug.utils import secure_filename
from fpdf import FPDF
from langchain_google_genai import ChatGoogleGenerativeAI
from langchain_core.prompts import PromptTemplate
from dotenv import load_dotenv

load_dotenv()

app = Flask(__name__)

app.config['ALLOWED_EXTENSIONS'] = {'pdf', 'txt', 'docx'}

# Initialize LLM
llm = ChatGoogleGenerativeAI(
    model="gemini-2.5-flash",
    api_key=os.getenv("GOOGLE_API_KEY"),
    temperature=0.0
)

# Prompt
mcq_prompt = PromptTemplate(
    input_variables=["context", "num_questions"],
    template="""
You are an AI assistant helping generate MCQs.

Text:
{context}

Generate {num_questions} MCQs.

Format strictly:
Question 1: ...
A) ...
B) ...
C) ...
D) ...
Correct Answer: A/B/C/D
"""
)

mcq_chain = mcq_prompt | llm

# -----------------------
# Helpers
# -----------------------

def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in app.config['ALLOWED_EXTENSIONS']


def extract_text_from_file(file, filename):
    ext = filename.rsplit('.', 1)[1].lower()

    if ext == 'pdf':
        with pdfplumber.open(io.BytesIO(file.read())) as pdf:
            return ''.join([p.extract_text() or '' for p in pdf.pages])

    elif ext == 'docx':
        doc = docx.Document(io.BytesIO(file.read()))
        return ' '.join([para.text for para in doc.paragraphs])

    elif ext == 'txt':
        return file.read().decode('utf-8')

    return None


def generate_mcqs(text, num_questions):
    response = mcq_chain.invoke({
        "context": text[:8000],  # prevent token overflow
        "num_questions": num_questions
    })
    return response.content.strip()


def parse_mcqs(mcq_text, quiz_title="Quiz"):
    questions = []
    blocks = re.split(r'Question \d+:', mcq_text)

    for block in blocks[1:]:
        lines = [l.strip() for l in block.strip().split('\n') if l.strip()]
        if len(lines) < 6:
            continue

        q = lines[0]
        options = []
        correct = -1

        for line in lines[1:]:
            if line.startswith(('A)', 'B)', 'C)', 'D)')):
                options.append(line[2:].strip())
            elif 'Correct Answer' in line:
                match = re.search(r'[A-D]', line.upper())
                if match:
                    correct = ord(match.group()) - ord('A')

        if len(options) == 4 and correct >= 0:
            questions.append({
                "question": q,
                "options": options,
                "correct_option": correct
            })

    return {"quiz_title": quiz_title, "questions": questions}


def create_pdf_buffer(mcqs_text):
    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Arial", size=12)

    for line in mcqs_text.split("\n"):
        pdf.multi_cell(0, 8, line)

    pdf_bytes = pdf.output(dest='S').encode('latin-1')
    return io.BytesIO(pdf_bytes)


# -----------------------
# Routes
# -----------------------

@app.route('/')
def index():
    return render_template('index.html')


@app.route('/generate', methods=['POST'])
def generate():
    if 'file' not in request.files:
        return "No file uploaded"

    file = request.files['file']
    if not (file and allowed_file(file.filename)):
        return "Invalid file"

    filename = secure_filename(file.filename)

    text = extract_text_from_file(file, filename)
    if not text:
        return "Failed to extract text"

    num_questions = int(request.form.get('num_questions', 5))

    mcqs = generate_mcqs(text, num_questions)

    return render_template(
        'results.html',
        mcqs=mcqs
    )


@app.route('/download/txt', methods=['POST'])
def download_txt():
    mcqs = request.form['mcqs']

    buffer = io.BytesIO(mcqs.encode('utf-8'))
    buffer.seek(0)

    return send_file(
        buffer,
        as_attachment=True,
        download_name="mcqs.txt",
        mimetype="text/plain"
    )


@app.route('/download/json', methods=['POST'])
def download_json():
    mcqs = request.form['mcqs']
    quiz_data = parse_mcqs(mcqs)

    buffer = io.BytesIO(json.dumps(quiz_data, indent=2).encode('utf-8'))
    buffer.seek(0)

    return send_file(
        buffer,
        as_attachment=True,
        download_name="mcqs.json",
        mimetype="application/json"
    )


@app.route('/download/pdf', methods=['POST'])
def download_pdf():
    mcqs = request.form['mcqs']
    buffer = create_pdf_buffer(mcqs)
    buffer.seek(0)

    return send_file(
        buffer,
        as_attachment=True,
        download_name="mcqs.pdf",
        mimetype="application/pdf"
    )


@app.route('/generate_json', methods=['POST'])
def generate_json():
    if 'file' not in request.files:
        return jsonify({"error": "No file uploaded"}), 400

    file = request.files['file']
    if not allowed_file(file.filename):
        return jsonify({"error": "Invalid file"}), 400

    filename = secure_filename(file.filename)
    text = extract_text_from_file(file, filename)

    if not text:
        return jsonify({"error": "Extraction failed"}), 400

    num_questions = int(request.form.get('num_questions', 5))

    mcqs = generate_mcqs(text, num_questions)
    quiz_data = parse_mcqs(mcqs)

    return jsonify(quiz_data)


# -----------------------
# Run
# -----------------------

if __name__ == "__main__":
    app.run(debug=False)
